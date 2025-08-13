# app.py - Advanced Resume Matcher with Enhanced Features

from flask import Flask, request, jsonify, render_template_string
from flask_cors import CORS
from werkzeug.utils import secure_filename
import os
import logging
import traceback
from datetime import datetime
import uuid
import json
from typing import Dict, List, Tuple, Optional, Any
from dataclasses import dataclass, asdict
import threading
from functools import wraps
import time
import concurrent.futures
from collections import Counter, defaultdict
import re
import hashlib
from pathlib import Path

# Document processing imports
try:
    import pdfplumber
except ImportError:
    pdfplumber = None
    print("Warning: pdfplumber not installed")

try:
    import docx
except ImportError:
    docx = None
    print("Warning: python-docx not installed")

try:
    from PIL import Image
    import pytesseract
except ImportError:
    pytesseract = None
    print("Warning: pytesseract or PIL not installed")

# ML and NLP imports
try:
    import spacy
    try:
        nlp = spacy.load("en_core_web_lg")
        print("Loaded en_core_web_lg model")
    except OSError:
        try:
            nlp = spacy.load("en_core_web_sm")
            print("Loaded en_core_web_sm model")
        except OSError:
            nlp = None
            print("Warning: No spaCy model found")
except ImportError:
    spacy = None
    nlp = None
    print("Warning: spaCy not installed")

try:
    import numpy as np
    from sklearn.feature_extraction.text import TfidfVectorizer
    from sklearn.metrics.pairwise import cosine_similarity
    from sklearn.cluster import KMeans
    from sklearn.decomposition import LatentDirichletAllocation
except ImportError:
    print("Warning: sklearn or numpy not installed")
    np = None
    TfidfVectorizer = None
    cosine_similarity = None

try:
    import nltk
    from nltk.tokenize import sent_tokenize, word_tokenize
    from nltk.corpus import stopwords
    from nltk.stem import WordNetLemmatizer
    from nltk.tag import pos_tag
    from nltk.chunk import ne_chunk
    # Download required data
    try:
        nltk.download('punkt', quiet=True)
        nltk.download('stopwords', quiet=True)
        nltk.download('averaged_perceptron_tagger', quiet=True)
        nltk.download('wordnet', quiet=True)
        nltk.download('maxent_ne_chunker', quiet=True)
        nltk.download('words', quiet=True)
    except:
        print("Warning: Could not download NLTK data")
except ImportError:
    nltk = None
    print("Warning: NLTK not installed")

# Advanced text processing
try:
    import textstat
    from textblob import TextBlob
    import yake
except ImportError:
    textstat = None
    TextBlob = None
    yake = None
    print("Warning: Advanced text processing libraries not installed")

# Initialize Flask app
app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 32 * 1024 * 1024  # 32MB max file size
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['CACHE_FOLDER'] = 'cache'
app.config['SECRET_KEY'] = 'your-secret-key-here'
CORS(app)

# Create directories
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['CACHE_FOLDER'], exist_ok=True)

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('resume_matcher.log'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# Enhanced data structures
@dataclass
class ContactInfo:
    """Contact information structure"""
    email: Optional[str] = None
    phone: Optional[str] = None
    linkedin: Optional[str] = None
    github: Optional[str] = None
    website: Optional[str] = None

@dataclass
class Experience:
    """Experience structure"""
    title: str
    company: str
    duration: str
    description: str
    years: float = 0.0

@dataclass
class Education:
    """Education structure"""
    degree: str
    institution: str
    year: Optional[str] = None
    gpa: Optional[str] = None

@dataclass
class ResumeAnalysis:
    """Comprehensive resume analysis"""
    contact_info: ContactInfo
    experiences: List[Experience]
    education: List[Education]
    skills: Dict[str, List[str]]
    certifications: List[str]
    total_experience: float
    seniority_level: str
    domain_expertise: List[str]
    readability_score: float
    keyword_density: Dict[str, float]
    ats_score: float

@dataclass
class JobAnalysis:
    """Job description analysis"""
    required_skills: Dict[str, List[str]]
    preferred_skills: Dict[str, List[str]]
    experience_level: str
    required_experience: float
    domain: str
    keywords: List[str]
    company_size: str
    role_type: str

@dataclass
class DetailedMatchingResult:
    """Enhanced matching result"""
    overall_score: float
    semantic_score: float
    keyword_score: float
    skills_score: float
    experience_score: float
    education_score: float
    ats_score: float
    cultural_fit_score: float
    
    matched_keywords: List[str]
    matched_skills: List[str]
    missing_critical_skills: List[str]
    skill_gaps: Dict[str, List[str]]
    
    recommendations: List[str]
    strengths: List[str]
    weaknesses: List[str]
    
    resume_analysis: ResumeAnalysis
    job_analysis: JobAnalysis
    
    processing_time: float
    confidence_score: float
    recommendation_priority: str

class CacheManager:
    """Simple cache manager for improved performance"""
    
    def __init__(self, cache_dir: str = 'cache'):
        self.cache_dir = Path(cache_dir)
        self.cache_dir.mkdir(exist_ok=True)
        self.cache_ttl = 3600  # 1 hour
    
    def get_cache_key(self, content: str) -> str:
        """Generate cache key from content"""
        return hashlib.md5(content.encode()).hexdigest()
    
    def get(self, key: str) -> Optional[Any]:
        """Get cached result"""
        try:
            cache_file = self.cache_dir / f"{key}.json"
            if cache_file.exists():
                stat = cache_file.stat()
                if time.time() - stat.st_mtime < self.cache_ttl:
                    with open(cache_file, 'r') as f:
                        return json.load(f)
        except Exception as e:
            logger.warning(f"Cache read error: {e}")
        return None
    
    def set(self, key: str, value: Any) -> None:
        """Set cached result"""
        try:
            cache_file = self.cache_dir / f"{key}.json"
            with open(cache_file, 'w') as f:
                json.dump(value, f, default=str)
        except Exception as e:
            logger.warning(f"Cache write error: {e}")

class AdvancedDocumentParser:
    """Advanced document parser with better extraction"""
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
        self.cache = CacheManager()
        
        # Advanced regex patterns
        self.patterns = {
            'email': re.compile(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'),
            'phone': re.compile(r'(\+?1[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'),
            'linkedin': re.compile(r'linkedin\.com/in/[^\s]+'),
            'github': re.compile(r'github\.com/[^\s]+'),
            'website': re.compile(r'https?://[^\s]+'),
            'experience': re.compile(r'(\d+)\s*(?:years?|yrs?)\s*(?:of\s*)?(?:experience|exp)', re.IGNORECASE),
            'gpa': re.compile(r'GPA:?\s*(\d+\.?\d*)', re.IGNORECASE)
        }
    
    def parse_document(self, file) -> Tuple[str, Dict[str, Any]]:
        """Enhanced document parsing with metadata extraction"""
        try:
            filename = secure_filename(file.filename)
            file_extension = os.path.splitext(filename)[1].lower()
            
            # Check cache first
            file_content = file.read()
            cache_key = self.cache.get_cache_key(f"{filename}_{len(file_content)}")
            cached_result = self.cache.get(cache_key)
            
            if cached_result:
                logger.info(f"Using cached result for {filename}")
                return cached_result['text'], cached_result['metadata']
            
            file.seek(0)
            
            file_info = {
                'filename': filename,
                'format': file_extension,
                'processed_at': datetime.now().isoformat(),
                'size': len(file_content),
                'processing_method': 'unknown'
            }
            
            # Parse based on file type
            if file_extension == '.pdf':
                text, metadata = self._parse_pdf_advanced(file)
                file_info['processing_method'] = 'pdf_advanced'
            elif file_extension in ['.docx', '.doc']:
                text, metadata = self._parse_docx_advanced(file)
                file_info['processing_method'] = 'docx_advanced'
            elif file_extension == '.txt':
                text = self._parse_txt(file_content)
                metadata = {}
                file_info['processing_method'] = 'txt_simple'
            else:
                text = self._fallback_text_extraction(file_content)
                metadata = {}
                file_info['processing_method'] = 'fallback'
            
            # Extract metadata
            file_info.update(metadata)
            file_info['character_count'] = len(text)
            file_info['word_count'] = len(text.split())
            file_info['paragraph_count'] = len(text.split('\n\n'))
            
            # Cache the result
            self.cache.set(cache_key, {'text': text, 'metadata': file_info})
            
            return text, file_info
            
        except Exception as e:
            self.logger.error(f"Error parsing document: {str(e)}")
            self.logger.error(traceback.format_exc())
            raise Exception(f"Could not parse document: {str(e)}")
    
    def _parse_pdf_advanced(self, file) -> Tuple[str, Dict[str, Any]]:
        """Advanced PDF parsing with metadata extraction"""
        try:
            text_parts = []
            metadata = {'pages': 0, 'images': 0, 'tables': 0}
            
            with pdfplumber.open(file) as pdf:
                metadata['pages'] = len(pdf.pages)
                
                for page in pdf.pages:
                    # Extract text
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                    
                    # Count images
                    if hasattr(page, 'images'):
                        metadata['images'] += len(page.images)
                    
                    # Count tables
                    tables = page.extract_tables()
                    if tables:
                        metadata['tables'] += len(tables)
            
            return '\n'.join(text_parts), metadata
            
        except Exception as e:
            self.logger.error(f"Advanced PDF parsing error: {str(e)}")
            # Fallback to simple extraction
            file.seek(0)
            return self._parse_pdf_simple(file), {}
    
    def _parse_pdf_simple(self, file) -> str:
        """Simple PDF parsing fallback"""
        try:
            text_parts = []
            with pdfplumber.open(file) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
            return '\n'.join(text_parts)
        except Exception as e:
            self.logger.error(f"Simple PDF parsing error: {str(e)}")
            raise
    
    def _parse_docx_advanced(self, file) -> Tuple[str, Dict[str, Any]]:
        """Advanced DOCX parsing with metadata extraction"""
        try:
            doc = docx.Document(file)
            text_parts = []
            metadata = {
                'paragraphs': 0,
                'tables': 0,
                'images': 0,
                'headers': 0,
                'footers': 0
            }
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
                    metadata['paragraphs'] += 1
            
            # Extract tables
            for table in doc.tables:
                metadata['tables'] += 1
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text)
                    text_parts.append('\t'.join(row_text))
            
            # Count images (approximate)
            try:
                for rel in doc.part.rels.values():
                    if "image" in rel.target_ref:
                        metadata['images'] += 1
            except:
                pass
            
            return '\n'.join(text_parts), metadata
            
        except Exception as e:
            self.logger.error(f"Advanced DOCX parsing error: {str(e)}")
            # Fallback to simple extraction
            file.seek(0)
            return self._parse_docx_simple(file), {}
    
    def _parse_docx_simple(self, file) -> str:
        """Simple DOCX parsing fallback"""
        try:
            doc = docx.Document(file)
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            return '\n'.join(text_parts)
        except Exception as e:
            self.logger.error(f"Simple DOCX parsing error: {str(e)}")
            raise
    
    def _parse_txt(self, file_content) -> str:
        """Enhanced TXT parsing with encoding detection"""
        encodings = ['utf-8', 'latin-1', 'ascii', 'utf-16', 'cp1252']
        
        for encoding in encodings:
            try:
                return file_content.decode(encoding)
            except UnicodeDecodeError:
                continue
        
        # Last resort
        return file_content.decode('utf-8', errors='ignore')
    
    def _fallback_text_extraction(self, file_content) -> str:
        """Enhanced fallback text extraction"""
        try:
            # Try different encodings
            return self._parse_txt(file_content)
        except:
            # Extract printable characters only
            return ''.join(chr(b) for b in file_content if 32 <= b <= 126)

class AdvancedResumeMatcher:
    """Advanced resume matcher with comprehensive analysis"""
    
    def __init__(self):
        self.nlp = nlp
        self.logger = logging.getLogger(__name__)
        self.cache = CacheManager()
        
        # Initialize lemmatizer
        self.lemmatizer = WordNetLemmatizer() if nltk else None
        
        # Enhanced skill categories with weights
        self.skill_categories = {
            'programming': {
                'skills': {
                    'python', 'java', 'javascript', 'typescript', 'c++', 'c#', 'ruby', 'go',
                    'php', 'html', 'css', 'sql', 'r', 'scala', 'kotlin', 'swift', 'rust',
                    'perl', 'shell', 'powershell', 'matlab', 'sas', 'dart', 'f#'
                },
                'weight': 0.25
            },
            'frameworks': {
                'skills': {
                    'react', 'angular', 'vue', 'django', 'flask', 'spring', 'express',
                    'rails', 'laravel', 'bootstrap', 'jquery', 'node.js', 'nodejs',
                    'asp.net', 'dotnet', '.net', 'symfony', 'codeigniter', 'fastapi',
                    'nextjs', 'nuxtjs', 'gatsby', 'svelte', 'ember'
                },
                'weight': 0.20
            },
            'databases': {
                'skills': {
                    'mysql', 'postgresql', 'mongodb', 'redis', 'sqlite', 'oracle',
                    'cassandra', 'elasticsearch', 'dynamodb', 'firebase', 'couchdb',
                    'mariadb', 'neo4j', 'influxdb', 'clickhouse'
                },
                'weight': 0.15
            },
            'cloud': {
                'skills': {
                    'aws', 'azure', 'gcp', 'google cloud', 'docker', 'kubernetes',
                    'terraform', 'ansible', 'jenkins', 'gitlab', 'github actions',
                    'circleci', 'travis', 'heroku', 'netlify', 'vercel'
                },
                'weight': 0.20
            },
            'tools': {
                'skills': {
                    'git', 'jira', 'confluence', 'slack', 'trello', 'asana',
                    'figma', 'sketch', 'photoshop', 'illustrator', 'tableau',
                    'power bi', 'excel', 'postman', 'swagger', 'kafka', 'rabbitmq'
                },
                'weight': 0.10
            },
            'methodologies': {
                'skills': {
                    'agile', 'scrum', 'kanban', 'devops', 'ci/cd', 'tdd', 'bdd',
                    'waterfall', 'lean', 'six sigma', 'itil', 'pmp'
                },
                'weight': 0.10
            }
        }
        
        # Industry domains
        self.domains = {
            'fintech': ['finance', 'banking', 'payment', 'trading', 'blockchain', 'cryptocurrency'],
            'healthcare': ['healthcare', 'medical', 'pharmaceutical', 'biotech', 'clinical'],
            'ecommerce': ['ecommerce', 'retail', 'shopping', 'marketplace', 'commerce'],
            'education': ['education', 'edtech', 'learning', 'training', 'academic'],
            'gaming': ['gaming', 'game', 'entertainment', 'mobile games', 'console'],
            'enterprise': ['enterprise', 'b2b', 'saas', 'crm', 'erp', 'business intelligence']
        }
        
        # Seniority levels
        self.seniority_patterns = {
            'entry': ['entry', 'junior', 'associate', 'trainee', 'intern', '0-2 years'],
            'mid': ['mid', 'middle', 'experienced', 'professional', '2-5 years', '3-7 years'],
            'senior': ['senior', 'lead', 'principal', 'expert', '5+ years', '7+ years'],
            'executive': ['manager', 'director', 'vp', 'cto', 'ceo', 'head', 'chief']
        }
    
    def extract_contact_info(self, text: str) -> ContactInfo:
        """Extract contact information from text"""
        contact = ContactInfo()
        
        # Email
        email_matches = re.findall(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', text)
        if email_matches:
            contact.email = email_matches[0]
        
        # Phone
        phone_matches = re.findall(r'(\+?1[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}', text)
        if phone_matches:
            contact.phone = phone_matches[0]
        
        # LinkedIn
        linkedin_matches = re.findall(r'linkedin\.com/in/[^\s]+', text, re.IGNORECASE)
        if linkedin_matches:
            contact.linkedin = linkedin_matches[0]
        
        # GitHub
        github_matches = re.findall(r'github\.com/[^\s]+', text, re.IGNORECASE)
        if github_matches:
            contact.github = github_matches[0]
        
        # Website
        website_matches = re.findall(r'https?://[^\s]+', text)
        if website_matches:
            contact.website = website_matches[0]
        
        return contact
    
    def extract_experiences(self, text: str) -> List[Experience]:
        """Extract work experiences from text"""
        experiences = []
        
        # This is a simplified version - in production, you'd use more sophisticated NLP
        experience_section = self._extract_section(text, ['experience', 'work history', 'employment'])
        
        if experience_section:
            # Split by common patterns
            entries = re.split(r'\n(?=[A-Z])', experience_section)
            
            for entry in entries:
                if len(entry.strip()) > 50:  # Minimum length for valid experience
                    lines = entry.strip().split('\n')
                    if len(lines) >= 2:
                        title = lines[0].strip()
                        company = lines[1].strip() if len(lines) > 1 else "Unknown"
                        duration = self._extract_duration(entry)
                        description = '\n'.join(lines[2:]) if len(lines) > 2 else ""
                        
                        years = self._calculate_years(duration)
                        
                        experiences.append(Experience(
                            title=title,
                            company=company,
                            duration=duration,
                            description=description,
                            years=years
                        ))
        
        return experiences
    
    def extract_education(self, text: str) -> List[Education]:
        """Extract education information from text"""
        education = []
        
        education_section = self._extract_section(text, ['education', 'academic', 'qualifications'])
        
        if education_section:
            # Common degree patterns
            degree_patterns = [
                r'(Bachelor|Master|PhD|MBA|MS|BS|BA|MA|MSc|BSc)[^\n]*',
                r'(B\.?[AES]\.?|M\.?[AES]\.?|Ph\.?D\.?)[^\n]*'
            ]
            
            for pattern in degree_patterns:
                matches = re.findall(pattern, education_section, re.IGNORECASE)
                for match in matches:
                    lines = match.split('\n')
                    degree = lines[0].strip()
                    institution = lines[1].strip() if len(lines) > 1 else "Unknown"
                    
                    # Extract year
                    year_match = re.search(r'20\d{2}|19\d{2}', match)
                    year = year_match.group() if year_match else None
                    
                    # Extract GPA
                    gpa_match = re.search(r'GPA:?\s*(\d+\.?\d*)', match, re.IGNORECASE)
                    gpa = gpa_match.group(1) if gpa_match else None
                    
                    education.append(Education(
                        degree=degree,
                        institution=institution,
                        year=year,
                        gpa=gpa
                    ))
        
        return education
    
    def extract_skills_advanced(self, text: str) -> Dict[str, List[str]]:
        """Advanced skill extraction with context analysis"""
        found_skills = defaultdict(list)
        text_lower = text.lower()
        
        # Extract skills by category
        for category, category_info in self.skill_categories.items():
            skills = category_info['skills']
            for skill in skills:
                # Use word boundaries for better matching
                pattern = r'\b' + re.escape(skill) + r'\b'
                if re.search(pattern, text_lower):
                    found_skills[category].append(skill)
        
        # Extract skills from dedicated sections
        skills_section = self._extract_section(text, ['skills', 'technical skills', 'technologies'])
        if skills_section:
            # Use YAKE for keyword extraction if available
            if yake:
                try:
                    kw_extractor = yake.KeywordExtractor(lan="en", n=3, dedupLim=0.7, top=20)
                    keywords = kw_extractor.extract_keywords(skills_section)
                    for score, keyword in keywords:
                        if score < 0.1:  # Low score means high relevance
                            found_skills['extracted'].append(keyword)
                except:
                    pass
        
        return dict(found_skills)
    
    def calculate_readability(self, text: str) -> float:
        """Calculate text readability score"""
        if not textstat:
            return 0.0
        
        try:
            # Use Flesch Reading Ease score
            return textstat.flesch_reading_ease(text)
        except:
            return 0.0
    
    def calculate_ats_score(self, text: str) -> float:
        """Calculate ATS (Applicant Tracking System) friendliness score"""
        score = 100.0
        
        # Check for common ATS issues
        if len(re.findall(r'[^\x00-\x7F]', text)) > 0:
            score -= 10  # Non-ASCII characters
        
        if len(re.findall(r'[|@#$%^&*]', text)) > 10:
            score -= 15  # Too many special characters
        
        if len(text.split()) < 100:
            score -= 20  # Too short
        
        # Check for standard sections
        standard_sections = ['experience', 'education', 'skills']
        found_sections = sum(1 for section in standard_sections if section in text.lower())
        if found_sections < 2:
            score -= 25
        
        return max(0, score)
    
    def analyze_resume(self, text: str) -> ResumeAnalysis:
        """Comprehensive resume analysis"""
        contact_info = self.extract_contact_info(text)
        experiences = self.extract_experiences(text)
        education = self.extract_education(text)
        skills = self.extract_skills_advanced(text)
        
        # Calculate total experience
        total_experience = sum(exp.years for exp in experiences)
        
        # Determine seniority level
        seniority_level = self._determine_seniority_level(text, total_experience)
        
        # Extract domain expertise
        domain_expertise = self._extract_domain_expertise(text)
        
        # Calculate readability
        readability_score = self.calculate_readability(text)
        
        # Calculate keyword density
        keyword_density = self._calculate_keyword_density(text)
        
        # Calculate ATS score
        ats_score = self.calculate_ats_score(text)
        
        # Extract certifications
        certifications = self._extract_certifications(text)
        
        return ResumeAnalysis(
            contact_info=contact_info,
            experiences=experiences,
            education=education,
            skills=skills,
            certifications=certifications,
            total_experience=total_experience,
            seniority_level=seniority_level,
            domain_expertise=domain_expertise,
            readability_score=readability_score,
            keyword_density=keyword_density,
            ats_score=ats_score
        )
    
    def analyze_job_description(self, text: str) -> JobAnalysis:
        """Analyze job description requirements"""
        # Extract required and preferred skills
        required_skills = self._extract_job_skills(text, ['required', 'must have', 'essential'])
        preferred_skills = self._extract_job_skills(text, ['preferred', 'nice to have', 'bonus'])
        
        # Determine experience level
        experience_level = self._determine_job_experience_level(text)
        
        # Extract required experience in years
        required_experience = self._extract_required_experience(text)
        
        # Determine domain
        domain = self._determine_job_domain(text)
        
        # Extract keywords
        keywords = self._extract_job_keywords(text)
        
        # Determine company size and role type
        company_size = self._determine_company_size(text)
        role_type = self._determine_role_type(text)
        
        return JobAnalysis(
            required_skills=required_skills,
            preferred_skills=preferred_skills,
            experience_level=experience_level,
            required_experience=required_experience,
            domain=domain,
            keywords=keywords,
            company_size=company_size,
            role_type=role_type
        )
    
    def calculate_cultural_fit(self, resume_text: str, job_text: str) -> float:
        """Calculate cultural fit score based on soft skills and values"""
        # Extract soft skills and values
        soft_skills = ['communication', 'teamwork', 'leadership', 'problem-solving', 
                      'creativity', 'adaptability', 'time management', 'critical thinking']
        
        resume_lower = resume_text.lower()
        job_lower = job_text.lower()
        
        matches = 0
        total_skills = len(soft_skills)
        
        for skill in soft_skills:
            if skill in resume_lower and skill in job_lower:
                matches += 1
        
        return (matches / total_skills) * 100 if total_skills > 0 else 0
    
    def generate_advanced_recommendations(self, resume_analysis: ResumeAnalysis, 
                                        job_analysis: JobAnalysis, 
                                        scores: Dict[str, float]) -> Tuple[List[str], List[str], List[str]]:
        """Generate detailed recommendations, strengths, and weaknesses"""
        recommendations = []
        strengths = []
        weaknesses = []
        
        # Skill gap analysis
        all_job_skills = set()
        for skills_list in job_analysis.required_skills.values():
            all_job_skills.update(skills_list)
        
        all_resume_skills = set()
        for skills_list in resume_analysis.skills.values():
            all_resume_skills.update(skills_list)
        
        missing_skills = all_job_skills - all_resume_skills
        
        # Generate recommendations
        if missing_skills:
            recommendations.append(f"Acquire these critical skills: {', '.join(list(missing_skills)[:5])}")
        
        if resume_analysis.total_experience < job_analysis.required_experience:
            recommendations.append(f"Gain more experience. You have {resume_analysis.total_experience} years, but {job_analysis.required_experience} years required.")
        
        if scores.get('ats_score', 0) < 70:
            recommendations.append("Improve ATS compatibility by using standard section headings and avoiding special characters.")
        
        if resume_analysis.readability_score < 50:
            recommendations.append("Improve resume readability by using simpler language and shorter sentences.")
        
        # Identify strengths
        if resume_analysis.total_experience >= job_analysis.required_experience:
            strengths.append(f"Strong experience background ({resume_analysis.total_experience} years)")
        
        if len(resume_analysis.skills) > 3:
            strengths.append("Diverse skill set across multiple categories")
        
        if resume_analysis.ats_score > 80:
            strengths.append("ATS-friendly resume format")
        
        # Identify weaknesses
        if not resume_analysis.contact_info.email:
            weaknesses.append("Missing contact information")
        
        if len(resume_analysis.education) == 0:
            weaknesses.append("No education information provided")
        
        if scores.get('skills_score', 0) < 40:
            weaknesses.append("Limited relevant technical skills")
        
        return recommendations, strengths, weaknesses
    
    def match_resume_advanced(self, resume_text: str, job_desc: str, file_info: Dict[str, Any]) -> DetailedMatchingResult:
        """Advanced resume matching with comprehensive analysis"""
        start_time = time.time()
        
        try:
            # Analyze resume and job description
            resume_analysis = self.analyze_resume(resume_text)
            job_analysis = self.analyze_job_description(job_desc)
            
            # Calculate various similarity scores
            semantic_score = self.semantic_similarity(resume_text, job_desc) * 100
            keyword_score, matched_keywords, missing_skills = self.keyword_matching(resume_text, job_desc)
            
            # Skills matching
            skills_score = self.calculate_skills_score(resume_analysis.skills, job_analysis.required_skills)
            
            # Experience matching
            experience_score = self.calculate_experience_score(resume_analysis.total_experience, job_analysis.required_experience)
            
            # Education matching
            education_score = self.calculate_education_score(resume_analysis.education, job_desc)
            
            # Cultural fit
            cultural_fit_score = self.calculate_cultural_fit(resume_text, job_desc)
            
            # Overall score calculation (weighted)
            weights = {
                'semantic': 0.20,
                'keyword': 0.25,
                'skills': 0.25,
                'experience': 0.15,
                'education': 0.10,
                'cultural_fit': 0.05
            }
            
            overall_score = (
                semantic_score * weights['semantic'] +
                keyword_score * weights['keyword'] +
                skills_score * weights['skills'] +
                experience_score * weights['experience'] +
                education_score * weights['education'] +
                cultural_fit_score * weights['cultural_fit']
            )
            
            # Calculate confidence score
            confidence_score = self.calculate_confidence_score(resume_analysis, job_analysis)
            
            # Generate detailed analysis
            recommendations, strengths, weaknesses = self.generate_advanced_recommendations(
                resume_analysis, job_analysis, {
                    'overall_score': overall_score,
                    'skills_score': skills_score,
                    'ats_score': resume_analysis.ats_score
                }
            )
            
            # Skill gap analysis
            skill_gaps = self.analyze_skill_gaps(resume_analysis.skills, job_analysis.required_skills)
            
            # Matched skills
            matched_skills = self.find_matched_skills(resume_analysis.skills, job_analysis.required_skills)
            
            # Determine priority
            recommendation_priority = self.determine_priority(overall_score, confidence_score)
            
            processing_time = time.time() - start_time
            
            return DetailedMatchingResult(
                overall_score=round(overall_score, 2),
                semantic_score=round(semantic_score, 2),
                keyword_score=round(keyword_score, 2),
                skills_score=round(skills_score, 2),
                experience_score=round(experience_score, 2),
                education_score=round(education_score, 2),
                ats_score=round(resume_analysis.ats_score, 2),
                cultural_fit_score=round(cultural_fit_score, 2),
                matched_keywords=matched_keywords,
                matched_skills=matched_skills,
                missing_critical_skills=missing_skills,
                skill_gaps=skill_gaps,
                recommendations=recommendations,
                strengths=strengths,
                weaknesses=weaknesses,
                resume_analysis=resume_analysis,
                job_analysis=job_analysis,
                processing_time=round(processing_time, 3),
                confidence_score=round(confidence_score, 2),
                recommendation_priority=recommendation_priority
            )
            
        except Exception as e:
            self.logger.error(f"Advanced resume matching failed: {str(e)}")
            self.logger.error(traceback.format_exc())
            raise
    
    # Helper methods (implementations would be here)
    def _extract_section(self, text: str, keywords: List[str]) -> str:
        """Extract specific section from text"""
        text_lower = text.lower()
        for keyword in keywords:
            pattern = rf'{keyword}.*?(?=\n[A-Z]|\n[A-Z]|\Z)'
            match = re.search(pattern, text_lower, re.DOTALL)
            if match:
                return match.group(0)
        return ""
    
    def _extract_duration(self, text: str) -> str:
        """Extract duration from text"""
        patterns = [
            r'(\d{4})\s*[-–]\s*(\d{4})',
            r'(\d{4})\s*[-–]\s*(present|current)',
            r'(\w+\s+\d{4})\s*[-–]\s*(\w+\s+\d{4})',
            r'(\w+\s+\d{4})\s*[-–]\s*(present|current)'
        ]
        
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                return match.group(0)
        return ""
    
    def _calculate_years(self, duration: str) -> float:
        """Calculate years from duration string"""
        if not duration:
            return 0.0
        
        # Extract years
        year_matches = re.findall(r'\d{4}', duration)
        if len(year_matches) >= 2:
            start_year = int(year_matches[0])
            end_year = int(year_matches[1]) if year_matches[1] else datetime.now().year
            return max(0, end_year - start_year)
        elif 'present' in duration.lower() or 'current' in duration.lower():
            if year_matches:
                start_year = int(year_matches[0])
                return max(0, datetime.now().year - start_year)
        
        return 0.0
    
    def _determine_seniority_level(self, text: str, total_experience: float) -> str:
        """Determine seniority level"""
        text_lower = text.lower()
        
        # Check for explicit seniority keywords
        for level, keywords in self.seniority_patterns.items():
            if any(keyword in text_lower for keyword in keywords):
                return level
        
        # Fallback to experience years
        if total_experience >= 10:
            return 'executive'
        elif total_experience >= 5:
            return 'senior'
        elif total_experience >= 2:
            return 'mid'
        else:
            return 'entry'
    
    def _extract_domain_expertise(self, text: str) -> List[str]:
        """Extract domain expertise"""
        text_lower = text.lower()
        domains = []
        
        for domain, keywords in self.domains.items():
            if any(keyword in text_lower for keyword in keywords):
                domains.append(domain)
        
        return domains
    
    def _calculate_keyword_density(self, text: str) -> Dict[str, float]:
        """Calculate keyword density"""
        words = text.lower().split()
        total_words = len(words)
        
        if total_words == 0:
            return {}
        
        word_counts = Counter(words)
        density = {}
        
        # Calculate density for top keywords
        for word, count in word_counts.most_common(10):
            if len(word) > 3:  # Only consider words longer than 3 characters
                density[word] = (count / total_words) * 100
        
        return density
    
    def _extract_certifications(self, text: str) -> List[str]:
        """Extract certifications"""
        cert_patterns = [
            r'(AWS|Azure|Google Cloud|GCP)\s+Certified',
            r'(PMP|CISSP|CISM|CISA|CEH)',
            r'(Certified|Certification).*?(?=\n|$)',
            r'(Oracle|Microsoft|Cisco|IBM)\s+Certified'
        ]
        
        certifications = []
        for pattern in cert_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            certifications.extend(matches)
        
        return list(set(certifications))
    
    def semantic_similarity(self, text1: str, text2: str) -> float:
        """Calculate semantic similarity using spaCy"""
        if not self.nlp:
            return self.basic_similarity(text1, text2)
        
        try:
            doc1 = self.nlp(text1[:1000000])  # Limit text length
            doc2 = self.nlp(text2[:1000000])
            
            if doc1.has_vector and doc2.has_vector:
                return doc1.similarity(doc2)
        except Exception as e:
            self.logger.warning(f"Semantic similarity failed: {str(e)}")
        
        return self.basic_similarity(text1, text2)
    
    def basic_similarity(self, text1: str, text2: str) -> float:
        """Basic similarity using word overlap"""
        words1 = set(text1.lower().split())
        words2 = set(text2.lower().split())
        
        if not words1 or not words2:
            return 0.0
        
        intersection = words1.intersection(words2)
        union = words1.union(words2)
        
        return len(intersection) / len(union) if union else 0.0
    
    def keyword_matching(self, resume_text: str, job_text: str) -> Tuple[float, List[str], List[str]]:
        """Enhanced keyword matching"""
        resume_words = set(resume_text.lower().split())
        job_words = set(job_text.lower().split())
        
        # Enhanced stop words
        stop_words = {
            'the', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'of', 'with', 'by',
            'is', 'are', 'was', 'were', 'be', 'been', 'being', 'have', 'has', 'had',
            'do', 'does', 'did', 'will', 'would', 'could', 'should', 'may', 'might',
            'must', 'can', 'this', 'that', 'these', 'those', 'a', 'an', 'we', 'you',
            'they', 'i', 'he', 'she', 'it', 'us', 'them', 'our', 'your', 'their'
        }
        
        resume_keywords = {word for word in resume_words if len(word) > 2 and word not in stop_words}
        job_keywords = {word for word in job_words if len(word) > 2 and word not in stop_words}
        
        matched = resume_keywords.intersection(job_keywords)
        
        score = (len(matched) / len(job_keywords) * 100) if job_keywords else 0
        
        # Find missing skills
        job_skills = self.extract_skills_advanced(job_text)
        resume_skills = self.extract_skills_advanced(resume_text)
        
        missing_skills = []
        for category, skills in job_skills.items():
            resume_category_skills = set(resume_skills.get(category, []))
            missing_skills.extend([skill for skill in skills if skill not in resume_category_skills])
        
        return score, list(matched)[:20], missing_skills[:10]
    
    def calculate_skills_score(self, resume_skills: Dict[str, List[str]], 
                             job_skills: Dict[str, List[str]]) -> float:
        """Calculate weighted skills score"""
        if not job_skills:
            return 0.0
        
        total_score = 0.0
        total_weight = 0.0
        
        for category, skills in job_skills.items():
            if category in self.skill_categories:
                weight = self.skill_categories[category]['weight']
                resume_category_skills = set(resume_skills.get(category, []))
                job_category_skills = set(skills)
                
                if job_category_skills:
                    category_score = len(resume_category_skills.intersection(job_category_skills)) / len(job_category_skills)
                    total_score += category_score * weight
                    total_weight += weight
        
        return (total_score / total_weight * 100) if total_weight > 0 else 0.0
    
    def calculate_experience_score(self, resume_experience: float, required_experience: float) -> float:
        """Calculate experience score"""
        if required_experience == 0:
            return 100.0
        
        if resume_experience >= required_experience:
            return 100.0
        else:
            return (resume_experience / required_experience) * 100
    
    def calculate_education_score(self, education: List[Education], job_desc: str) -> float:
        """Calculate education score"""
        if not education:
            return 0.0
        
        job_lower = job_desc.lower()
        
        # Check for degree requirements
        if any(degree in job_lower for degree in ['bachelor', 'master', 'phd', 'mba']):
            has_degree = any(degree.degree for degree in education)
            return 100.0 if has_degree else 50.0
        
        return 80.0  # Default score if no specific education requirements
    
    def calculate_confidence_score(self, resume_analysis: ResumeAnalysis, 
                                 job_analysis: JobAnalysis) -> float:
        """Calculate confidence score for the match"""
        factors = []
        
        # Text quality factors
        if resume_analysis.readability_score > 60:
            factors.append(20)
        
        if resume_analysis.ats_score > 70:
            factors.append(15)
        
        # Content completeness
        if resume_analysis.contact_info.email:
            factors.append(10)
        
        if len(resume_analysis.experiences) > 0:
            factors.append(15)
        
        if len(resume_analysis.education) > 0:
            factors.append(10)
        
        if len(resume_analysis.skills) > 2:
            factors.append(15)
        
        # Job description quality
        if len(job_analysis.keywords) > 10:
            factors.append(15)
        
        return sum(factors)
    
    def determine_priority(self, overall_score: float, confidence_score: float) -> str:
        """Determine recommendation priority"""
        if overall_score >= 80 and confidence_score >= 70:
            return "High"
        elif overall_score >= 60 and confidence_score >= 50:
            return "Medium"
        else:
            return "Low"
    
    def analyze_skill_gaps(self, resume_skills: Dict[str, List[str]], 
                          job_skills: Dict[str, List[str]]) -> Dict[str, List[str]]:
        """Analyze skill gaps by category"""
        gaps = {}
        
        for category, skills in job_skills.items():
            resume_category_skills = set(resume_skills.get(category, []))
            job_category_skills = set(skills)
            missing = job_category_skills - resume_category_skills
            
            if missing:
                gaps[category] = list(missing)
        
        return gaps
    
    def find_matched_skills(self, resume_skills: Dict[str, List[str]], 
                           job_skills: Dict[str, List[str]]) -> List[str]:
        """Find matched skills across all categories"""
        matched = []
        
        for category, skills in job_skills.items():
            resume_category_skills = set(resume_skills.get(category, []))
            job_category_skills = set(skills)
            matched.extend(resume_category_skills.intersection(job_category_skills))
        
        return list(set(matched))
    
    # Additional helper methods for job analysis
    def _extract_job_skills(self, text: str, indicators: List[str]) -> Dict[str, List[str]]:
        """Extract job skills based on indicators"""
        skills = defaultdict(list)
        
        for indicator in indicators:
            # Find sections with the indicator
            pattern = rf'{indicator}.*?(?=\n[A-Z]|\n[A-Z]|\Z)'
            matches = re.findall(pattern, text, re.IGNORECASE | re.DOTALL)
            
            for match in matches:
                # Extract skills from the match
                for category, category_info in self.skill_categories.items():
                    category_skills = category_info['skills']
                    for skill in category_skills:
                        if skill in match.lower():
                            skills[category].append(skill)
        
        return dict(skills)
    
    def _determine_job_experience_level(self, text: str) -> str:
        """Determine job experience level"""
        text_lower = text.lower()
        
        for level, keywords in self.seniority_patterns.items():
            if any(keyword in text_lower for keyword in keywords):
                return level
        
        return 'mid'  # Default
    
    def _extract_required_experience(self, text: str) -> float:
        """Extract required experience in years"""
        patterns = [
            r'(\d+)\+?\s*(?:years?|yrs?)\s*(?:of\s*)?(?:experience|exp)',
            r'(\d+)\s*to\s*(\d+)\s*(?:years?|yrs?)',
            r'minimum\s*(\d+)\s*(?:years?|yrs?)'
        ]
        
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                return float(match.group(1))
        
        return 0.0
    
    def _determine_job_domain(self, text: str) -> str:
        """Determine job domain"""
        text_lower = text.lower()
        
        for domain, keywords in self.domains.items():
            if any(keyword in text_lower for keyword in keywords):
                return domain
        
        return 'general'
    
    def _extract_job_keywords(self, text: str) -> List[str]:
        """Extract important keywords from job description"""
        # This would use more sophisticated NLP in production
        words = text.lower().split()
        
        # Filter common words and get important ones
        stop_words = {
            'the', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'of', 'with', 'by',
            'is', 'are', 'was', 'were', 'be', 'been', 'being', 'have', 'has', 'had'
        }
        
        keywords = [word for word in words if len(word) > 3 and word not in stop_words]
        
        return list(set(keywords))
    
    def _determine_company_size(self, text: str) -> str:
        """Determine company size from job description"""
        text_lower = text.lower()
        
        if any(term in text_lower for term in ['startup', 'small company', 'growing team']):
            return 'small'
        elif any(term in text_lower for term in ['enterprise', 'fortune', 'multinational']):
            return 'large'
        else:
            return 'medium'
    
    def _determine_role_type(self, text: str) -> str:
        """Determine role type"""
        text_lower = text.lower()
        
        if any(term in text_lower for term in ['remote', 'work from home', 'distributed']):
            return 'remote'
        elif any(term in text_lower for term in ['hybrid', 'flexible']):
            return 'hybrid'
        else:
            return 'onsite'

# Initialize components
document_parser = AdvancedDocumentParser()
resume_matcher = AdvancedResumeMatcher()

# Enhanced HTML template with better UI
ADVANCED_HTML_TEMPLATE = """
<!DOCTYPE html>
<html>
<head>
    <title>Advanced Resume Matcher</title>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <style>
        * {
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }
        
        body {
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            min-height: 100vh;
            padding: 20px;
        }
        
        .container {
            max-width: 1200px;
            margin: 0 auto;
            background: rgba(255, 255, 255, 0.95);
            border-radius: 20px;
            padding: 30px;
            box-shadow: 0 20px 40px rgba(0, 0, 0, 0.1);
        }
        
        .header {
            text-align: center;
            margin-bottom: 30px;
        }
        
        .header h1 {
            color: #333;
            font-size: 2.5em;
            margin-bottom: 10px;
            background: linear-gradient(45deg, #667eea, #764ba2);
            -webkit-background-clip: text;
            -webkit-text-fill-color: transparent;
            background-clip: text;
        }
        
        .header p {
            color: #666;
            font-size: 1.2em;
        }
        
        .form-section {
            background: white;
            padding: 25px;
            border-radius: 15px;
            margin-bottom: 20px;
            box-shadow: 0 5px 15px rgba(0, 0, 0, 0.08);
        }
        
        .form-group {
            margin-bottom: 20px;
        }
        
        label {
            display: block;
            margin-bottom: 8px;
            font-weight: 600;
            color: #333;
        }
        
        input[type="file"] {
            width: 100%;
            padding: 12px;
            border: 2px dashed #ddd;
            border-radius: 10px;
            background: #f9f9f9;
            transition: border-color 0.3s;
        }
        
        input[type="file"]:hover {
            border-color: #667eea;
        }
        
        textarea {
            width: 100%;
            padding: 15px;
            border: 2px solid #ddd;
            border-radius: 10px;
            font-size: 14px;
            resize: vertical;
            transition: border-color 0.3s;
        }
        
        textarea:focus {
            outline: none;
            border-color: #667eea;
        }
        
        .btn-primary {
            background: linear-gradient(45deg, #667eea, #764ba2);
            color: white;
            padding: 15px 30px;
            border: none;
            border-radius: 10px;
            font-size: 16px;
            font-weight: 600;
            cursor: pointer;
            transition: transform 0.3s;
        }
        
        .btn-primary:hover {
            transform: translateY(-2px);
            box-shadow: 0 5px 15px rgba(102, 126, 234, 0.3);
        }
        
        .loading {
            text-align: center;
            padding: 40px;
            display: none;
        }
        
        .loading-spinner {
            border: 4px solid #f3f3f3;
            border-top: 4px solid #667eea;
            border-radius: 50%;
            width: 40px;
            height: 40px;
            animation: spin 1s linear infinite;
            margin: 0 auto 20px;
        }
        
        @keyframes spin {
            0% { transform: rotate(0deg); }
            100% { transform: rotate(360deg); }
        }
        
        .error {
            color: #dc3545;
            background: #ffe6e6;
            padding: 15px;
            border-radius: 10px;
            margin: 15px 0;
            border: 1px solid #dc3545;
        }
        
        .results {
            display: none;
            margin-top: 20px;
        }
        
        .score-grid {
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(250px, 1fr));
            gap: 20px;
            margin-bottom: 30px;
        }
        
        .score-card {
            background: white;
            padding: 20px;
            border-radius: 15px;
            box-shadow: 0 5px 15px rgba(0, 0, 0, 0.1);
            text-align: center;
            transition: transform 0.3s;
        }
        
        .score-card:hover {
            transform: translateY(-5px);
        }
        
        .score-value {
            font-size: 2.5em;
            font-weight: bold;
            margin-bottom: 10px;
        }
        
        .score-label {
            color: #666;
            font-size: 14px;
        }
        
        .overall-score .score-value {
            color: #28a745;
        }
        
        .semantic-score .score-value {
            color: #007bff;
        }
        
        .skills-score .score-value {
            color: #ffc107;
        }
        
        .ats-score .score-value {
            color: #17a2b8;
        }
        
        .details-section {
            background: white;
            padding: 25px;
            border-radius: 15px;
            margin-bottom: 20px;
            box-shadow: 0 5px 15px rgba(0, 0, 0, 0.08);
        }
        
        .details-section h3 {
            color: #333;
            margin-bottom: 15px;
            font-size: 1.3em;
        }
        
        .skill-tags {
            display: flex;
            flex-wrap: wrap;
            gap: 10px;
            margin-bottom: 15px;
        }
        
        .skill-tag {
            background: #667eea;
            color: white;
            padding: 6px 12px;
            border-radius: 20px;
            font-size: 12px;
            font-weight: 500;
        }
        
        .missing-skill-tag {
            background: #dc3545;
        }
        
        .recommendation-item {
            background: #f8f9fa;
            padding: 15px;
            border-radius: 10px;
            margin-bottom: 10px;
            border-left: 4px solid #667eea;
        }
        
        .strength-item {
            background: #d4edda;
            border-left-color: #28a745;
        }
        
        .weakness-item {
            background: #f8d7da;
            border-left-color: #dc3545;
        }
        
        .tabs {
            display: flex;
            border-bottom: 2px solid #ddd;
            margin-bottom: 20px;
        }
        
        .tab {
            padding: 10px 20px;
            cursor: pointer;
            background: #f8f9fa;
            border: none;
            border-bottom: 2px solid transparent;
            transition: all 0.3s;
        }
        
        .tab.active {
            background: #667eea;
            color: white;
            border-bottom-color: #667eea;
        }
        
        .tab-content {
            display: none;
        }
        
        .tab-content.active {
            display: block;
        }
        
        .priority-high {
            border-left-color: #dc3545;
        }
        
        .priority-medium {
            border-left-color: #ffc107;
        }
        
        .priority-low {
            border-left-color: #28a745;
        }
        
        @media (max-width: 768px) {
            .container {
                padding: 20px;
            }
            
            .header h1 {
                font-size: 2em;
            }
            
            .score-grid {
                grid-template-columns: 1fr;
            }
        }
    </style>
</head>
<body>
    <div class="container">
        <div class="header">
            <h1>🚀 Advanced Resume Matcher</h1>
            <p>AI-powered resume analysis and job matching</p>
        </div>
        
        <div class="form-section">
            <form id="uploadForm" enctype="multipart/form-data">
                <div class="form-group">
                    <label for="resume">📄 Upload Resume (PDF, DOCX, TXT)</label>
                    <input type="file" id="resume" name="resume" accept=".pdf,.docx,.txt" required>
                </div>
                
                <div class="form-group">
                    <label for="job_description">📋 Job Description</label>
                    <textarea id="job_description" name="job_description" rows="10" 
                              placeholder="Paste the job description here..." required></textarea>
                </div>
                
                <button type="submit" class="btn-primary">🔍 Analyze Resume</button>
            </form>
        </div>
        
        <div class="loading" id="loading">
            <div class="loading-spinner"></div>
            <p>Analyzing your resume with advanced AI...</p>
        </div>
        
        <div class="error" id="error" style="display: none;"></div>
        
        <div class="results" id="results">
            <div class="score-grid" id="scoreGrid">
                <!-- Score cards will be populated here -->
            </div>
            
            <div class="tabs">
                <button class="tab active" onclick="showTab('overview')">📊 Overview</button>
                <button class="tab" onclick="showTab('skills')">🛠️ Skills</button>
                <button class="tab" onclick="showTab('recommendations')">💡 Recommendations</button>
                <button class="tab" onclick="showTab('analysis')">📈 Analysis</button>
            </div>
            
            <div id="overview" class="tab-content active">
                <div class="details-section">
                    <h3>📋 Match Summary</h3>
                    <div id="matchSummary"></div>
                </div>
            </div>
            
            <div id="skills" class="tab-content">
                <div class="details-section">
                    <h3>✅ Matched Skills</h3>
                    <div id="matchedSkills"></div>
                </div>
                
                <div class="details-section">
                    <h3>❌ Missing Skills</h3>
                    <div id="missingSkills"></div>
                </div>
            </div>
            
            <div id="recommendations" class="tab-content">
                <div class="details-section">
                    <h3>💪 Strengths</h3>
                    <div id="strengths"></div>
                </div>
                
                <div class="details-section">
                    <h3>⚠️ Areas for Improvement</h3>
                    <div id="weaknesses"></div>
                </div>
                
                <div class="details-section">
                    <h3>🎯 Recommendations</h3>
                    <div id="recommendationsList"></div>
                </div>
            </div>
            
            <div id="analysis" class="tab-content">
                <div class="details-section">
                    <h3>📊 Detailed Analysis</h3>
                    <div id="detailedAnalysis"></div>
                </div>
            </div>
        </div>
    </div>

    <script>
        let currentResults = null;
        
        document.getElementById('uploadForm').addEventListener('submit', async function(e) {
            e.preventDefault();
            
            const formData = new FormData();
            const resumeFile = document.getElementById('resume').files[0];
            const jobDescription = document.getElementById('job_description').value;
            
            if (!resumeFile) {
                showError('Please select a resume file');
                return;
            }
            
            if (!jobDescription.trim()) {
                showError('Please enter a job description');
                return;
            }
            
            formData.append('resume', resumeFile);
            formData.append('job_description', jobDescription);
            
            showLoading(true);
            hideError();
            hideResults();
            
            try {
                const response = await fetch('/upload', {
                    method: 'POST',
                    body: formData
                });
                
                const result = await response.json();
                
                if (response.ok) {
                    currentResults = result;
                    displayResults(result);
                } else {
                    showError(result.error || 'An error occurred');
                }
            } catch (error) {
                showError('Network error: ' + error.message);
            } finally {
                showLoading(false);
            }
        });
        
        function showLoading(show) {
            document.getElementById('loading').style.display = show ? 'block' : 'none';
        }
        
        function showError(message) {
            document.getElementById('error').textContent = message;
            document.getElementById('error').style.display = 'block';
        }
        
        function hideError() {
            document.getElementById('error').style.display = 'none';
        }
        
        function hideResults() {
            document.getElementById('results').style.display = 'none';
        }
        
        function displayResults(data) {
            // Display score cards
            const scoreGrid = document.getElementById('scoreGrid');
            scoreGrid.innerHTML = `
                <div class="score-card overall-score">
                    <div class="score-value">${data.overall_score}%</div>
                    <div class="score-label">Overall Match</div>
                </div>
                <div class="score-card semantic-score">
                    <div class="score-value">${data.semantic_score}%</div>
                    <div class="score-label">Semantic Match</div>
                </div>
                <div class="score-card skills-score">
                    <div class="score-value">${data.skills_score}%</div>
                    <div class="score-label">Skills Match</div>
                </div>
                <div class="score-card ats-score">
                    <div class="score-value">${data.ats_score}%</div>
                    <div class="score-label">ATS Score</div>
                </div>
                <div class="score-card">
                    <div class="score-value">${data.experience_score}%</div>
                    <div class="score-label">Experience Match</div>
                </div>
                <div class="score-card">
                    <div class="score-value">${data.cultural_fit_score}%</div>
                    <div class="score-label">Cultural Fit</div>
                </div>
            `;
            
            // Display match summary
            const matchSummary = document.getElementById('matchSummary');
            matchSummary.innerHTML = `
                <div class="recommendation-item priority-${data.recommendation_priority.toLowerCase()}">
                    <strong>Priority:</strong> ${data.recommendation_priority}<br>
                    <strong>Confidence:</strong> ${data.confidence_score}%<br>
                    <strong>Processing Time:</strong> ${data.processing_time}s
                </div>
                <p><strong>Resume Analysis:</strong> ${data.resume_analysis.seniority_level} level candidate with ${data.resume_analysis.total_experience} years of experience</p>
                <p><strong>Domain Match:</strong> ${data.resume_analysis.domain_expertise.join(', ') || 'General'}</p>
            `;
            
            // Display skills
            displaySkills(data);
            
            // Display recommendations
            displayRecommendations(data);
            
            // Display detailed analysis
            displayDetailedAnalysis(data);
            
            document.getElementById('results').style.display = 'block';
        }
        
        function displaySkills(data) {
            const matchedSkills = document.getElementById('matchedSkills');
            const missingSkills = document.getElementById('missingSkills');
            
            if (data.matched_skills.length > 0) {
                matchedSkills.innerHTML = '<div class="skill-tags">' + 
                    data.matched_skills.map(skill => `<span class="skill-tag">${skill}</span>`).join('') + 
                    '</div>';
            } else {
                matchedSkills.innerHTML = '<p>No matched skills found</p>';
            }
            
            if (data.missing_critical_skills.length > 0) {
                missingSkills.innerHTML = '<div class="skill-tags">' + 
                    data.missing_critical_skills.map(skill => `<span class="skill-tag missing-skill-tag">${skill}</span>`).join('') + 
                    '</div>';
            } else {
                missingSkills.innerHTML = '<p>No critical skills missing</p>';
            }
        }
        
        function displayRecommendations(data) {
            const strengths = document.getElementById('strengths');
            const weaknesses = document.getElementById('weaknesses');
            const recommendations = document.getElementById('recommendationsList');
            
            if (data.strengths.length > 0) {
                strengths.innerHTML = data.strengths.map(strength => 
                    `<div class="recommendation-item strength-item">${strength}</div>`
                ).join('');
            } else {
                strengths.innerHTML = '<p>No specific strengths identified</p>';
            }
            
            if (data.weaknesses.length > 0) {
                weaknesses.innerHTML = data.weaknesses.map(weakness => 
                    `<div class="recommendation-item weakness-item">${weakness}</div>`
                ).join('');
            } else {
                weaknesses.innerHTML = '<p>No specific weaknesses identified</p>';
            }
            
            if (data.recommendations.length > 0) {
                recommendations.innerHTML = data.recommendations.map(rec => 
                    `<div class="recommendation-item">${rec}</div>`
                ).join('');
            } else {
                recommendations.innerHTML = '<p>No specific recommendations</p>';
            }
        }
        
        function displayDetailedAnalysis(data) {
            const detailedAnalysis = document.getElementById('detailedAnalysis');
            
            detailedAnalysis.innerHTML = `
                <h4>Resume Quality Metrics</h4>
                <p><strong>Readability Score:</strong> ${data.resume_analysis.readability_score.toFixed(1)}</p>
                <p><strong>Contact Info:</strong> ${data.resume_analysis.contact_info.email ? '✅ Complete' : '❌ Incomplete'}</p>
                <p><strong>Experiences:</strong> ${data.resume_analysis.experiences.length} positions</p>
                <p><strong>Education:</strong> ${data.resume_analysis.education.length} entries</p>
                <p><strong>Certifications:</strong> ${data.resume_analysis.certifications.length}</p>
                
                <h4>Job Analysis</h4>
                <p><strong>Required Experience:</strong> ${data.job_analysis.required_experience} years</p>
                <p><strong>Experience Level:</strong> ${data.job_analysis.experience_level}</p>
                <p><strong>Domain:</strong> ${data.job_analysis.domain}</p>
                <p><strong>Role Type:</strong> ${data.job_analysis.role_type}</p>
                
                <h4>Skill Gap Analysis</h4>
                <div id="skillGaps"></div>
            `;
            
            // Display skill gaps
            const skillGaps = document.getElementById('skillGaps');
            if (data.skill_gaps && Object.keys(data.skill_gaps).length > 0) {
                let gapsHtml = '';
                for (const [category, skills] of Object.entries(data.skill_gaps)) {
                    gapsHtml += `<p><strong>${category}:</strong> ${skills.join(', ')}</p>`;
                }
                skillGaps.innerHTML = gapsHtml;
            } else {
                skillGaps.innerHTML = '<p>No significant skill gaps identified</p>';
            }
        }
        
        function showTab(tabName) {
            // Hide all tab contents
            const tabContents = document.querySelectorAll('.tab-content');
            tabContents.forEach(content => content.classList.remove('active'));
            
            // Remove active class from all tabs
            const tabs = document.querySelectorAll('.tab');
            tabs.forEach(tab => tab.classList.remove('active'));
            
            // Show selected tab content
            document.getElementById(tabName).classList.add('active');
            
            // Add active class to selected tab
            event.target.classList.add('active');
        }
        
        // File input enhancement
        document.getElementById('resume').addEventListener('change', function(e) {
            const file = e.target.files[0];
            if (file) {
                const fileSize = (file.size / 1024 / 1024).toFixed(2);
                console.log(`File selected: ${file.name} (${fileSize} MB)`);
            }
        });
    </script>
</body>
</html>
"""

# Enhanced Flask Routes
@app.route('/')
def index():
    return render_template_string(ADVANCED_HTML_TEMPLATE)

@app.route('/upload', methods=['POST'])
def upload_resume():
    """Enhanced resume upload with advanced analysis"""
    try:
        logger.info("Advanced upload request received")
        
        # Validation
        if 'resume' not in request.files:
            logger.error("No resume file in request")
            return jsonify({'error': 'No resume file provided'}), 400
        
        if 'job_description' not in request.form:
            logger.error("No job description in request")
            return jsonify({'error': 'No job description provided'}), 400
        
        file = request.files['resume']
        job_desc = request.form['job_description'].strip()
        
        if file.filename == '':
            logger.error("No file selected")
            return jsonify({'error': 'No file selected'}), 400
        
        if not job_desc or len(job_desc) < 50:
            logger.error("Insufficient job description")
            return jsonify({'error': 'Job description must be at least 50 characters'}), 400
        
        logger.info(f"Processing file: {file.filename}")
        
        # Parse document with advanced parser
        try:
            resume_text, file_info = document_parser.parse_document(file)
            logger.info(f"Document parsed successfully. Text length: {len(resume_text)}")
        except Exception as e:
            logger.error(f"Document parsing failed: {str(e)}")
            return jsonify({'error': f'Could not parse document: {str(e)}'}), 400
        
        if not resume_text or len(resume_text.strip()) < 50:
            logger.error("Insufficient text extracted")
            return jsonify({'error': 'Could not extract sufficient text from the resume'}), 400
        
        # Advanced resume matching
        try:
            result = resume_matcher.match_resume_advanced(resume_text, job_desc, file_info)
            logger.info(f"Advanced resume matching completed. Score: {result.overall_score}%")
        except Exception as e:
            logger.error(f"Advanced resume matching failed: {str(e)}")
            return jsonify({'error': f'Analysis failed: {str(e)}'}), 500
        
        # Convert nested dataclasses to dict
        def convert_dataclass(obj):
            if hasattr(obj, '__dict__'):
                return {k: convert_dataclass(v) for k, v in obj.__dict__.items()}
            elif isinstance(obj, list):
                return [convert_dataclass(item) for item in obj]
            elif isinstance(obj, dict):
                return {k: convert_dataclass(v) for k, v in obj.items()}
            else:
                return obj
        
        response_data = convert_dataclass(result)
        
        return jsonify(response_data)
        
    except Exception as e:
        logger.error(f"Unexpected error in upload_resume: {str(e)}")
        logger.error(traceback.format_exc())
        return jsonify({'error': f'Internal server error: {str(e)}'}), 500

@app.route('/health', methods=['GET'])
def health_check():
    """Enhanced health check endpoint"""
    return jsonify({
        'status': 'healthy',
        'version': '2.0',
        'features': {
            'spacy_available': nlp is not None,
            'pdfplumber_available': pdfplumber is not None,
            'docx_available': docx is not None,
            'advanced_nlp': textstat is not None,
            'caching_enabled': True,
            'concurrent_processing': True
        },
        'timestamp': datetime.now().isoformat()
    })

@app.route('/api/batch-analyze', methods=['POST'])
def batch_analyze():
    """Batch analysis endpoint for multiple resumes"""
    try:
        if 'resumes' not in request.files:
            return jsonify({'error': 'No resume files provided'}), 400
        
        files = request.files.getlist('resumes')
        job_desc = request.form.get('job_description', '').strip()
        
        if not job_desc:
            return jsonify({'error': 'Job description required'}), 400
        
        results = []
        
        # Process files concurrently
        with concurrent.futures.ThreadPoolExecutor(max_workers=3) as executor:
            futures = []
            
            for file in files:
                if file.filename:
                    future = executor.submit(process_single_resume, file, job_desc)
                    futures.append((file.filename, future))
            
            for filename, future in futures:
                try:
                    result = future.result(timeout=30)  # 30-second timeout
                    results.append({
                        'filename': filename,
                        'success': True,
                        'result': result
                    })
                except Exception as e:
                    results.append({
                        'filename': filename,
                        'success': False,
                        'error': str(e)
                    })
        
        return jsonify({
            'processed': len(results),
            'results': results
        })
        
    except Exception as e:
        logger.error(f"Batch analysis failed: {str(e)}")
        return jsonify({'error': str(e)}), 500

def process_single_resume(file, job_desc):
    """Process a single resume file"""
    try:
        resume_text, file_info = document_parser.parse_document(file)
        result = resume_matcher.match_resume_advanced(resume_text, job_desc, file_info)
        return asdict(result)
    except Exception as e:
        raise Exception(f"Processing failed: {str(e)}")

if __name__ == '__main__':
    logger.info("Starting Advanced Resume Matcher Application")
    logger.info("=" * 50)
    logger.info("🚀 Advanced Features Enabled:")
    logger.info(f"  📊 SpaCy NLP: {'✅' if nlp else '❌'}")
    logger.info(f"  📄 PDF Support: {'✅' if pdfplumber else '❌'}")
    logger.info(f"  📝 DOCX Support: {'✅' if docx else '❌'}")
    logger.info(f"  🔍 Advanced Text Analysis: {'✅' if textstat else '❌'}")
    logger.info(f"  💾 Caching: ✅")
    logger.info(f"  🔄 Concurrent Processing: ✅")
    logger.info(f"  🎯 Batch Processing: ✅")
    logger.info("=" * 50)
    
    app.run(host='0.0.0.0', port=5000, debug=True)
